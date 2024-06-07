import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from docx import Document
from docx.shared import Inches

# Cargar los datos desde el archivo Excel
file_path = 'datos.xlsx'  # Reemplaza con el nombre correcto del archivo
df = pd.read_excel(file_path, sheet_name='Datos')

# Crear un nuevo documento Word
doc = Document()
doc.add_heading('Informe de Análisis de Datos', 0)

# Clasificación de Variables
doc.add_heading('1. Clasificación de Variables', level=1)
classification_text = (
    "Estado: Cualitativa Nominal\n"
    "Sexo: Cualitativa Nominal\n"
    "Edad: Cuantitativa Continua\n"
    "Raza: Cualitativa Nominal\n"
    "Años de Estudio: Cuantitativa Continua\n"
    "Ingresos: Cuantitativa Continua\n"
    "Altura: Cuantitativa Continua"
)
doc.add_paragraph(classification_text)

# Tabla de Distribución de Frecuencias
doc.add_heading('2. Tabla de Distribución de Frecuencias', level=1)

# Estado
doc.add_heading('Estado', level=2)
estado_freq = df['Estado'].value_counts()
estado_table = doc.add_table(rows=1, cols=2)
hdr_cells = estado_table.rows[0].cells
hdr_cells[0].text = 'Estado'
hdr_cells[1].text = 'Frecuencia'
for var, freq in estado_freq.items():
    row_cells = estado_table.add_row().cells
    row_cells[0].text = str(var)
    row_cells[1].text = str(freq)

# Años de Estudio
doc.add_heading('Años de estudio', level=2)
años_estudio_freq = df['Años de estudio'].value_counts()
años_estudio_table = doc.add_table(rows=1, cols=2)
hdr_cells = años_estudio_table.rows[0].cells
hdr_cells[0].text = 'Años de estudio'
hdr_cells[1].text = 'Frecuencia'
for var, freq in años_estudio_freq.items():
    row_cells = años_estudio_table.add_row().cells
    row_cells[0].text = str(var)
    row_cells[1].text = str(freq)

# Representación Gráfica
doc.add_heading('3. Representación Gráfica', level=1)

# Sexo
doc.add_heading('Sexo', level=2)
plt.figure(figsize=(10, 6))
sns.countplot(x='Sexo', data=df)
plt.title('Distribución de Sexo')
plt.savefig('sexo.png')
doc.add_picture('sexo.png', width=Inches(5))
plt.close()

# Edad
doc.add_heading('Edad', level=2)
plt.figure(figsize=(10, 6))
df['Edad'].hist()
plt.title('Distribución de Edad')
plt.xlabel('Edad')
plt.ylabel('Frecuencia')
plt.savefig('edad.png')
doc.add_picture('edad.png', width=Inches(5))
plt.close()

# Raza
doc.add_heading('Color', level=2)
plt.figure(figsize=(10, 6))
sns.countplot(x='Color', data=df)
plt.title('Distribución de Color')
plt.savefig('Color.png')
doc.add_picture('Color.png', width=Inches(5))
plt.close()

# Altura
doc.add_heading('Altura', level=2)
plt.figure(figsize=(10, 6))
df['Altura'].hist()
plt.title('Distribución de Altura')
plt.xlabel('Altura')
plt.ylabel('Frecuencia')
plt.savefig('altura.png')
doc.add_picture('altura.png', width=Inches(5))
plt.close()

# Tabla con Medidas Resumen
doc.add_heading('4. Tabla con Medidas Resumen', level=1)
summary_stats = df.describe()
summary_table = doc.add_table(rows=summary_stats.shape[0]+1, cols=summary_stats.shape[1]+1)
hdr_cells = summary_table.rows[0].cells
hdr_cells[0].text = 'Medida'
for i, column in enumerate(summary_stats.columns):
    hdr_cells[i+1].text = column

for i, index in enumerate(summary_stats.index):
    row_cells = summary_table.add_row().cells
    row_cells[0].text = index
    for j, column in enumerate(summary_stats.columns):
        row_cells[j+1].text = str(round(summary_stats.loc[index, column], 2))

# Gráficos de Distribución de Frecuencias
doc.add_heading('5. Gráficos de Distribución de Frecuencias', level=1)

# Ingreso según Sexo
doc.add_heading('Ingresos según Sexo', level=2)
plt.figure(figsize=(10, 6))
sns.boxplot(x='Sexo', y='Ingresos', data=df)
plt.title('Ingresos según Sexo')
plt.savefig('ingresos_sexo.png')
doc.add_picture('ingresos_sexo.png', width=Inches(5))
plt.close()

# Altura según Raza
doc.add_heading('Altura según Raza', level=2)
plt.figure(figsize=(10, 6))
sns.boxplot(x='Color', y='Altura', data=df)
plt.title('Altura según Raza')
plt.savefig('altura_raza.png')
doc.add_picture('altura_raza.png', width=Inches(5))
plt.close()

# Ingreso según Raza
doc.add_heading('Ingresos según Raza', level=2)
plt.figure(figsize=(10, 6))
sns.boxplot(x='Color', y='Ingresos', data=df)
plt.title('Ingresos según Raza')
plt.savefig('ingresos_raza.png')
doc.add_picture('ingresos_raza.png', width=Inches(5))
plt.close()

# Gráfico de Asociación
doc.add_heading('6. Gráfico de Asociación', level=1)
plt.figure(figsize=(10, 6))
sns.scatterplot(x='Años de estudio', y='Ingresos', data=df)
plt.title('Ingresos según Años de estudio')
plt.savefig('asociacion.png')
doc.add_picture('asociacion.png', width=Inches(5))
plt.close()

# Categorizar la Variable Ingreso
doc.add_heading('7. Categorizar la Variable Ingresos', level=1)
bins = [0, 2*788, 4*788, 10*788, 20*788, float('inf')]
labels = ['Menor a 2 SM', 'Mayor o igual a 2 SM y menor a 4 SM', 'Mayor o igual a 4 SM y menor a 10 SM', 'Mayor o igual a 10 SM y menor a 20 SM', 'Mayor o igual a 20 SM']
df['Ingresos Categoría'] = pd.cut(df['Ingresos'], bins=bins, labels=labels)

plt.figure(figsize=(10, 6))
sns.countplot(x='Ingresos Categoría', data=df)
plt.title('Distribución de Ingresos Categoría')
plt.savefig('ingresos_categoria.png')
doc.add_picture('ingresos_categoria.png', width=Inches(5))
plt.close()

# Intervalo de Confianza
doc.add_heading('8. Intervalo de Confianza', level=1)
ingresos = df['Ingresos'].dropna()
mean_ingresos = ingresos.mean()
confidence_interval = stats.t.interval(0.95, len(ingresos)-1, loc=mean_ingresos, scale=stats.sem(ingresos))
confidence_text = f"Intervalo de confianza del 95% para la media de ingresos: {confidence_interval}"
doc.add_paragraph(confidence_text)

# Guardar el documento
doc.save('Actividad_Grupo_numero_del_grupo.docx')

print("El informe ha sido generado y guardado como 'Actividad_Grupo_numero_del_grupo.docx'.")
